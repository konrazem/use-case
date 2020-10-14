from classes.use_case import UseCase
import openpyxl
from docx import Document
from docx.shared import Inches # need for layout in px

docx_name = 'Use Case - description.docx'
xlsx_name = 'data.xlsx'

wb = openpyxl.load_workbook(xlsx_name)

categories = ['ACCOUNT MANAGEMENT', 'BANK ACCOUNT MANAGEMENT', 'TRANSACTION MANAGEMENT', 'SECURITY', 'NOTIFICATION MANAGEMENT', 'HELP', 'FEEDBACK', 'COMMUNITY']



# -----------------------------------------------------
# Create Document
# -----------------------------------------------------

document = Document()
document.add_heading('Use Case Descriptions', 0)
appendix = []

# -----------------------------------------------------

category_number = 1

for category in categories:

    sheet = UseCase(wb[category])
    _data = sheet.getTables()
    tables = _data['tables']
    appendix.append(_data['appendix'].copy())


    # ADD HEADING
    document.add_heading(str(category_number) + '. ' +
                         categories[category_number - 1],
                         level=1)

    # ADD PICTURE
    document.add_picture('images/' + category + '.png', width=Inches(5.0))
    # ADD SPACE
    document.add_paragraph('\n')
    # ADD TABLES
    table_number = 1
    for table in tables:

        document.add_heading(str(category_number) + '.' + str(table_number) + '. ' +
                             table['Name'] + ' (' + str(table['Priority']) +
                             ')',
                             level=2)

        # table var convert to table :)

        _table = document.add_table(rows=len(table), cols=2)
        _table.autofit = False
        _table.allow_autofit = False
        # set table col width
        _table.columns[0].width = Inches(1.0)
        _table.columns[1].width = Inches(5.0)

        i = 0
        for key in table:
            _table.rows[i].cells[0].paragraphs[0].add_run(key).bold = True
            _table.rows[i].cells[0].width = Inches(1.0)
            _table.rows[i].cells[1].text = table[key]
            _table.rows[i].cells[1].width = Inches(5.0)
            i = i + 1

        # ADD SPACE
        document.add_paragraph('\n')
        table_number = table_number + 1

    category_number = category_number + 1
    document.add_page_break()

    # -----------------------------------------------------
    # Create Appendix
    # -----------------------------------------------------



# ADD HEADING
document.add_heading( 'APPENDIX', level=1)

# ADD TABLES
atable_number = 1
for i in appendix:
    for atable in i:
        # ADD HEADING
        document.add_heading('A.' + str(atable_number) +
                             '. ' + atable['Name'] + ' (' +
                             str(atable['Priority']) + ')',
                             level=2)

        # ADD TABLE
        _table = document.add_table(rows=len(atable), cols=2)
        _table.autofit = False
        _table.allow_autofit = False
        # set table col width
        _table.columns[0].width = Inches(1.0)
        _table.columns[1].width = Inches(5.0)

        i = 0
        for key in atable:
            _table.rows[i].cells[0].paragraphs[0].add_run(key).bold = True
            _table.rows[i].cells[0].width = Inches(1.0)
            _table.rows[i].cells[1].text = atable[key]
            _table.rows[i].cells[1].width = Inches(5.0)
            i = i + 1


        atable_number = atable_number + 1
        # ADD SPACE
        document.add_paragraph('\n')



# -----------------------------------------------------
document.save(docx_name)
print('Saved in: ')
print(docx_name)